from pathlib import Path
import sys

sys.path.insert(0, "/private/tmp/chapel-docx")

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/manuals/OLV_Ceremonies_Manual_Working_Draft.docx")

NAVY = RGBColor(31, 77, 120)
BLUE = RGBColor(46, 116, 181)
PALE_BLUE = "DDEBF7"
PALE_GREEN = "E2F0D9"
PALE_GOLD = "FFF2CC"
PALE_GRAY = "F2F2F2"
DARK = RGBColor(35, 35, 35)
GRAY = RGBColor(95, 95, 95)
WHITE = RGBColor(255, 255, 255)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_run(run, size=10.5, bold=False, italic=False, color=DARK, font="Aptos"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr_text, fld_sep, text, fld_end])


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color, before, after in (
        ("Heading 1", 18, BLUE, 12, 7),
        ("Heading 2", 13.5, NAVY, 9, 4),
        ("Heading 3", 11.5, NAVY, 7, 3),
    ):
        style = styles[name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Aptos"
        style.font.size = Pt(10.25)
        style.paragraph_format.space_after = Pt(2)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(header.add_run("OLV OPERATIONS PLATFORM  |  CEREMONIES MANUAL"), size=8.5, bold=True, color=GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(footer.add_run("Working draft  •  July 16, 2026  •  Page "), size=8.5, color=GRAY)
    add_field(footer, "PAGE")


def add_para(doc, text="", bold_lead=None, italic=False):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        set_run(p.add_run(bold_lead), bold=True, color=NAVY)
        set_run(p.add_run(text[len(bold_lead):]), italic=italic)
    else:
        set_run(p.add_run(text), italic=italic)
    return p


def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        set_run(p.add_run(item), size=10.25)


def add_numbered(doc, items):
    for index, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.34)
        p.paragraph_format.first_line_indent = Inches(-0.34)
        set_run(p.add_run(f"{index}."), size=10.25)
        set_run(p.add_run(f"\t{item}"), size=10.25)


def add_status_box(doc, label, text, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run(f"{label}: "), bold=True, color=NAVY)
    set_run(p.add_run(text))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, value in enumerate(headers):
        set_cell_shading(header.cells[idx], NAVY.__str__().replace("RGBColor(", "").replace(")", "").replace(", ", ""))
        p = header.cells[idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(value), size=9.5, bold=True, color=WHITE)
    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        if row_idx % 2:
            for cell in cells:
                set_cell_shading(cell, PALE_GRAY)
        for idx, value in enumerate(values):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run(p.add_run(str(value)), size=9.25, bold=(idx == 0), color=NAVY if idx == 0 else DARK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


doc = Document()
configure_document(doc)

# Cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(80)
set_run(p.add_run("OUR LADY OF VICTORY CHAPEL"), size=11, bold=True, color=BLUE)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
set_run(p.add_run("Ceremonies Manual"), size=32, bold=True, color=DARK, font="Aptos Display")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(22)
set_run(p.add_run("Operational planning, preparation, staffing, and communication"), size=16, color=GRAY, font="Aptos Display")

add_status_box(doc, "DOCUMENT STATUS", "Working draft compiled from approved product workflows. Stakeholder and clergy review remains required for items specifically marked for review.", PALE_GOLD)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(18)
set_run(p.add_run("Purpose"), size=13, bold=True, color=NAVY)
add_para(doc, "Give clergy, assistants, ceremony coordinators, ministry leaders, and technical stewards one practical reference for how OLV ceremonies are scheduled, prepared, staffed, communicated, and reviewed.")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
set_run(p.add_run("Source hierarchy"), size=13, bold=True, color=NAVY)
add_bullets(doc, [
    "Father's liturgical judgment and authorized direction",
    "The SSPX 1962 Ordo and authentic liturgical books",
    "Approved OLV workflow specifications and effective-dated local rules",
    "This manual as an operator-oriented summary",
])

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(60)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run(p.add_run("OLV Operations Platform"), size=12, bold=True, color=BLUE)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run(p.add_run("Working Draft  •  July 16, 2026"), size=10.5, color=GRAY)

doc.add_page_break()

# How to use
doc.add_heading("How to use this manual", level=1)
add_para(doc, "This manual organizes real chapel operations rather than software screens. It is not a substitute for the Roman Missal, Roman Ritual, SSPX Ordo, ministry training, or Father's direction.")

add_table(doc, ["Label", "Meaning"], [
    ("Approved Practice", "Accepted by the Product Owner for product design and ready for stakeholder validation or implementation planning."),
    ("Stakeholder Review", "A known decision remains open; do not silently invent a local rule."),
    ("Future Development", "The platform must allow the capability later, but it is not required for the initial Scheduling release."),
], [2200, 7160])

doc.add_heading("Human-Centered Ministry Principle", level=2)
add_status_box(doc, "GUIDING PRINCIPLE", "The platform supports the ministries of the chapel; it does not replace them. It should automate repetitive administration, improve communication, and reduce errors while preserving meaningful opportunities for parishioners to contribute time, talents, and judgment.", PALE_GREEN)

doc.add_heading("Authority and stewardship", level=2)
add_table(doc, ["Role", "Responsibility"], [
    ("Father or authorized delegate", "Final liturgical authority; confirms exceptions, controlling directions, and pastoral decisions."),
    ("Main Chapel Administrator", "Controls service dates, times, locations, recurring schedules, publication, cancellations, and facilities."),
    ("Ceremony Coordinator", "Coordinates liturgical plans, ceremony templates, preparation, and cross-ministry implications."),
    ("Father's Assistant", "Handles family-facing intake, calendar holds, communication, and decisions communicated by Father."),
    ("Workflow Steward", "Defines, refines, and approves a functional workflow within delegated authority."),
    ("Technical Steward", "Evaluates reusable components, integrations, security, and implementation options."),
    ("Ministry Leader", "Confirms ministry requirements, qualifications, staffing, readiness, and exceptions."),
], [2700, 6660])

add_para(doc, "Public documentation uses role titles rather than personal names. Current assignments and delegated authority belong in a private project register.")

doc.add_heading("Cross-functional cases", level=2)
add_para(doc, "A major ceremony or event may link several workstreams without collapsing them into one task. A Confirmation, major feast, or procession may include liturgy, Schola, hospitality, food, facilities, venue reservation, communications, fundraising, setup, cleanup, and volunteer staffing.")
add_bullets(doc, [
    "Each workstream retains its own owner, permissions, tasks, status, and acceptance criteria.",
    "The overall case shares dates, dependencies, resource conflicts, and change notifications.",
    "Unreviewed areas are labeled SME input required rather than filled by assumption.",
    "The Webmaster is the designated Schola Workflow Steward and Technical Steward; existing Schola components are evaluated only after the desired workflow is defined.",
])

# Calendar
doc.add_heading("1. Calendar and service foundations", level=1)
add_status_box(doc, "APPROVED PRACTICE", "The SSPX 1962 Ordo is the preferred liturgical source. The platform advises on rubrics but never replaces Father's judgment.", PALE_GREEN)

doc.add_heading("Calendar horizons", level=2)
add_bullets(doc, [
    "Liturgical reference calendar: view as far ahead as Ordo data permits.",
    "Operational chapel calendar: generate actual services and staffing within a rolling two-month horizon.",
    "Public reference data may show feast or feria, class, color, and commemorations.",
    "Gloria, Credo, preface, Mass options, staffing, and private notes remain role-restricted.",
])

doc.add_heading("Ordinary schedule", level=2)
add_table(doc, ["Day", "Masses", "Confessions and Rosary"], [
    ("Sunday", "7:00, 9:00, 11:00 a.m.; 5:00 p.m. — 11:00 normally Sung", "30 minutes before each Mass"),
    ("Monday", "7:15 a.m.", "One hour before"),
    ("Friday", "6:30 p.m.", "One hour before"),
    ("Saturday", "9:00 a.m.", "One hour before"),
], [1800, 4460, 3100])

doc.add_heading("SSPX proper calendar", level=2)
add_bullets(doc, [
    "Import the SSPX proper calendar, not only the general 1960 calendar.",
    "All SSPX-specific first-class feasts default to Sung Mass.",
    "Our Lady of Compassion and St. Pius X are minimum validation cases.",
    "External Solemnity eligibility and precedence follow the rubrics; Father confirms uncertain cases.",
    "If an import conflicts with a known SSPX rank, retain the last verified rank and alert the Ceremony Coordinator.",
])

doc.add_heading("First Friday, First Saturday, and summer", level=2)
add_bullets(doc, [
    "First Saturday governs overnight Adoration; First Friday alone does not.",
    "Friday before First Saturday: 6:30 p.m. Mass, 7:00 p.m. Exposition, Adoration 7:30 p.m.–7:30 a.m.",
    "First Saturday: Devotions and Benediction 8:00–9:00 a.m.; Mass at 9:00 a.m.",
    "Sacred Heart and Immaculate Heart preferences apply when rubrically permitted.",
    "Summer Low-Mass schedule begins after the External Solemnity of Saints Peter and Paul and ends with the Assumption; August 15 is Sung and following Sundays return to Sung Mass.",
])

doc.add_heading("Holy Days and Christmas", level=2)
add_table(doc, ["Occurrence", "Schedule"], [
    ("Weekday Holy Day", "7:00 a.m. Low Mass; 7:00 p.m. Sung Mass"),
    ("Saturday Holy Day", "9:00 a.m. Sung Mass only"),
    ("Sunday Holy Day", "Normal Sunday schedule; 11:00 a.m. Sung"),
    ("Christmas", "12:00 a.m. Sung Mass; 9:00 and 11:00 a.m. Low Masses; replaces ordinary schedule"),
], [2500, 6860])

# Holy Week
doc.add_heading("2. Holy Week and Easter", level=1)
add_table(doc, ["Event", "Time and operational rule"], [
    ("Palm Sunday", "11:00 a.m. start; palms and procession precede Mass."),
    ("Holy Thursday Mass", "7:00 p.m."),
    ("Holy Thursday Adoration", "8:30 p.m.–midnight."),
    ("Good Friday Rosary", "1:45 p.m.; independently staffed through usher ministry."),
    ("Solemn Stations", "2:15 p.m.; Cross Bearer, Ac1, Ac2; no sacristan assignment."),
    ("Solemn Liturgy", "3:00 p.m.; independently staffed."),
    ("Easter Vigil with Baptisms", "10:00 p.m. start; midnight Mass; one combined assignment."),
    ("Easter Vigil without Baptisms", "10:30 p.m. start; midnight Mass; one combined assignment."),
    ("Easter Sunday", "9:00 and 11:00 a.m. Low Masses; normal Sunday Confessions and Rosary."),
], [2800, 6560])

doc.add_heading("Holy Thursday Mandatum", level=2)
add_bullets(doc, [
    "The Holy Name Society President invites twelve men as Apostles.",
    "An Apostle may not also serve at the altar.",
    "An Apostle may also serve as sacristan.",
    "An Apostle may continue ushering when the usher ministry leader approves the dual assignment.",
])

add_status_box(doc, "FUTURE DEVELOPMENT", "Tenebrae is not currently generated but may be added as an annual template.", PALE_GOLD)

# Processions
doc.add_heading("3. Processions, patronal feasts, and Foundation Masses", level=1)
add_table(doc, ["Observance", "Ceremony rule"], [
    ("Corpus Christi", "External Solemnity at following Sunday's 11:00 a.m. Sung Mass; procession afterward."),
    ("Christ the King", "Procession after the 11:00 a.m. Sung Mass."),
    ("Crowning of Our Lady", "First Sunday in May; after the 11:00 a.m. Sung Mass."),
    ("Our Lady of Victory", "First Sunday in October; procession after the 11:00 a.m. Sung Mass."),
    ("St. Philomena — January 10", "Dated Foundation Mass remains on January 10; Low unless Sunday. Add External Solemnity and procession on eligible Sunday."),
    ("January 10 on Sunday", "Celebrate directly as 11:00 a.m. Sung patronal Mass with procession; do not label External Solemnity."),
], [3100, 6260])

doc.add_heading("St. Philomena Foundation Masses", level=2)
add_table(doc, ["Date", "Observance"], [
    ("January 10", "Birthday"),
    ("May 25", "Finding of her relics"),
    ("August 10", "Martyrdom and translation of her body"),
    ("August 11", "Patronal feast"),
], [2200, 7160])

doc.add_heading("Invitation-only bearers", level=2)
add_bullets(doc, [
    "The Holy Name Society President invites four men for every procession.",
    "Corpus Christi and Christ the King use Canopy Bearers.",
    "Marian and St. Philomena processions use Bier Bearers.",
    "Invitees may differ between processions.",
])

doc.add_heading("Banner Bearers", level=2)
add_table(doc, ["Banner", "Positions"], [
    ("Holy Name Society", "2"),
    ("Cor Unum", "1"),
    ("OLV", "1 — omitted for Corpus Christi and Christ the King"),
    ("Lepanto", "1"),
    ("Sacred Heart", "1"),
    ("Christ the King", "1"),
], [3200, 6160])
add_para(doc, "Banner Bearers are event volunteers coordinated and approved by the Holy Name Society President.")

# Blessings
doc.add_heading("4. Annual blessings and transferred local practices", level=1)
add_table(doc, ["Rite", "Current operational rule"], [
    ("Distribution of Ashes", "No automatic Ash Wednesday Mass under current mission coverage. Distribute ashes after every Mass on the following Sunday; start times unchanged."),
    ("Candlemas", "Full blessing, distribution, procession, and Mass. Saturday 9:00 a.m.; Sunday 11:00 a.m. Weekday creates review item for Father."),
    ("Blessing of Throats", "February 3 when weekday Mass is available; also after every Mass on following Sunday. If Feb. 3 is Sunday, that Sunday only."),
    ("Epiphany Water", "Separate January 5 rite; proper date preserved. Alternate timing requires Father's authorization."),
    ("Blessing of Chalk", "Separate January 6 rite; proper date preserved. Alternate timing requires Father's authorization."),
    ("Blessing of Wine", "Separate event after Mass on December 27; nearest priest-covered opportunity may be proposed for approval."),
    ("Blessing of Arms", "Restricted sacristy event after Mass on September 29 or following Saturday; restricted visibility."),
], [2600, 6760])

add_status_box(doc, "STAKEHOLDER REVIEW", "Confirm whether Epiphany water and chalk are fixed annual events or ad libitum, confirm exact times, and inventory any additional annual blessings currently practiced at OLV.", PALE_GOLD)

# Wedding
doc.add_heading("5. Weddings", level=1)
add_status_box(doc, "WORKFLOW STATUS", "Approved for stakeholder review; authentic ceremony-form terminology and exact OLV forms remain to be validated by Father and stakeholders.", PALE_GOLD)

doc.add_heading("Operating sequence", level=2)
add_numbered(doc, [
    "Father or Father's Assistant receives the request; the assistant creates a restricted case and tentative chapel hold.",
    "The system checks chapel and priest conflicts. A conflict may be recorded but blocks confirmation.",
    "Father or expected celebrant approves directly or through recorded verbal delegation.",
    "The assistant records the ceremony form and optional linked rehearsal.",
    "Draft staffing is generated for sacristans, servers, ushers, and Schola when Sung.",
    "Ministry leaders review staffing; family information and internal preparation remain separate.",
    "Publication is explicit and requires separate permission for names.",
])

doc.add_heading("Ceremony and preparation rules", level=2)
add_bullets(doc, [
    "Provisional forms: Low Nuptial Mass, Sung Nuptial Mass, and Marriage Ceremony without Mass.",
    "The rehearsal is a restricted linked event that reserves the chapel and notifies priest and assistant; it creates no ordinary staffing unless added.",
    "Sacramental-record, diocesan-permission, and related statuses may be tracked without storing the underlying protected records.",
    "Material date, time, form, or priest changes require renewed approval.",
    "The case remains pending until explicitly approved or cancelled; it does not expire silently.",
])

# Funeral
doc.add_heading("6. Funerals", level=1)
add_status_box(doc, "PASTORAL STANDARD", "Use compassionate, light-touch intake. Capture information once and route only what each role needs; do not make the family navigate liturgical terminology.", PALE_GREEN)

doc.add_heading("Compassionate first contact", level=2)
add_bullets(doc, [
    "Initially require only proposed date, time, location, case type, and restricted family contact.",
    "Allow save-and-stop without treating missing details as errors.",
    "Place a restricted tentative hold and check chapel and priest conflicts.",
    "After Father or delegate approves the date, mark Confirmed—details pending.",
])

doc.add_heading("Family conversation guide", level=2)
add_bullets(doc, [
    "Ask only what the family can reasonably answer: body present, viewing, Rosary, Rosary leader, burial, related activities, communication preference, and publication permission.",
    "Support Unknown/TBD, conditional questions, automatic saving, phone-first use, and continuation across contacts.",
    "Send a warm, editable summary for family verification; requested changes require human review.",
])

doc.add_heading("Internal liturgical plan", level=2)
add_para(doc, "The Ceremony Coordinator develops the internal plan with Father. Possible linked events include Viewing, Rosary, Reception of the Body, Requiem Mass, Absolution at the coffin, Absolution at the catafalque, procession to burial, and burial service. The system may suggest; it does not decide.")

doc.add_heading("Funeral timing and Schola", level=2)
add_bullets(doc, [
    "Rosary defaults to thirty minutes before the funeral.",
    "When viewing and Rosary both occur, suggest viewing 60–30 minutes before and Rosary 30–0 minutes before.",
    "If OLV leads the Rosary, create a Lead Rosary usher duty; a funeral-party leader's name is optional.",
    "When the body is present, the plan may request the Subvenite. If Schola is unavailable, Father decides.",
])

doc.add_heading("Funeral server templates", level=2)
add_table(doc, ["Form", "Staffing rule"], [
    ("Low Requiem without body/catafalque", "1 server required; second optional."),
    ("Low Requiem with body/catafalque", "Prefer MC, Th, Ac1, Ac2, Cross Bearer; 1 acolyte absolute minimum."),
    ("Sung Requiem", "MC, Th, Ac1, Ac2, Cross Bearer, and Schola required unless Father authorizes exception."),
], [3300, 6060])

doc.add_heading("Communication integrity", level=2)
add_bullets(doc, [
    "Record Father's verbal direction with source, wording, author, time, and resulting changes.",
    "Notify every affected ministry; do not assume one recipient will relay it.",
    "Conflicting reports remain visible until Father, his assistant, or the Ceremony Coordinator identifies the controlling direction.",
])

# Baptism
doc.add_heading("7. Baptisms", level=1)
add_numbered(doc, [
    "Father or the assistant creates the request; the assistant records candidates, approximate age when needed, proposed schedule, language, contact, and reception request.",
    "Check chapel, facilities, priest calendar, and priest-language compatibility.",
    "A conflict permits a tentative hold but blocks confirmation; a language mismatch warns but may be overridden.",
    "The assistant confirms after conflicts clear and the system creates one required sacristan position.",
    "The assistant sends a reviewed family summary and publishes only specifically authorized information.",
])

doc.add_heading("Baptism rules", level=2)
add_bullets(doc, [
    "Candidate classification: Infant, Adult, or Unsure—enter age. Age supports preparation but never decides the rite automatically.",
    "One case may contain multiple family members. Duration defaults to one hour plus ten minutes for each additional candidate.",
    "Ordinary template: one sacristan, no server. A future template may add a server.",
    "A reception reserves the hall with default access one hour before and after; the family handles setup and cleanup unless separate duties are created.",
    "Easter Vigil Baptism is a linked case within the existing Vigil event and updates duration, scope, checklist, and ministry notices.",
])

add_status_box(doc, "STAKEHOLDER REVIEW", "Confirm godparent/preparation status boundaries, future server needs, final retention policy, and detailed Easter Vigil effects.", PALE_GOLD)

# Requiem selection
doc.add_heading("8. Selecting a Requiem for a scheduled Mass", level=1)
add_para(doc, "This is a fast operational change to an existing Mass, not a Funeral case. Do not collect donor, offering, deceased-person, or intention details.")

add_numbered(doc, [
    "The priest or assistant selects Requiem on the scheduled Mass.",
    "Low or Sung is required; the occasion/formulary is optional.",
    "The system displays a discreet rubrical advisory without blocking Father.",
    "Notify the assigned sacristan immediately; Sung selections also notify server and Schola leaders.",
    "Apply the staffing template, preserve compatible assignments, and send incompatible assignments for leader review.",
])

doc.add_heading("Optional formulary guide", level=2)
add_table(doc, ["Occasion", "Suggested formulary"], [
    ("Funeral or Exequial Mass", "In die obitus seu depositionis defuncti"),
    ("Day of death, burial, news of death, or reburial", "Pro die obitus"),
    ("Third, seventh, or thirtieth day", "Pro die obitus with applicable prayers"),
    ("Anniversary", "In anniversario defunctorum"),
    ("Other Mass for the dead", "In Missis quotidianis defunctorum"),
], [4100, 5260])
add_para(doc, "If the optional selection is blank, the sacristan sees Requiem formulary not specified. No reminder is generated.")

doc.add_heading("Last-minute coordination", level=2)
add_table(doc, ["Notice before Mass", "Acknowledgment rule"], [
    ("1–3 hours", "15 minutes to acknowledge"),
    ("30–60 minutes", "5 minutes to acknowledge"),
    ("Under 30 minutes", "Notify assigned personnel, leaders, and alternates together"),
], [2800, 6560])
add_para(doc, "Require one Ready or Cannot staff response per affected ministry. Automated reminders stop thirty minutes before Mass; Father and his assistant receive one final readiness summary.")

# Privacy and communications
doc.add_heading("9. Communication, privacy, and readiness", level=1)
add_bullets(doc, [
    "Public calendars show services, not volunteer names, staffing counts, private notes, family contacts, or restricted events.",
    "Family contact data is stored separately from calendar and ministry views.",
    "Tentative and confirmed notices must be visually and textually distinct.",
    "Material changes use a review-and-notify preview so affected roles receive the same controlling information.",
    "Public cancellation notices are neutral; restricted reasons remain visible only to authorized leaders.",
    "Preferred and backup channels are used according to urgency and delivery status.",
    "Readiness may be tracked independently for liturgy, sacristy, servers, ushers, and Schola.",
])

doc.add_heading("Handling conflicting verbal information", level=2)
add_numbered(doc, [
    "Record each report without overwriting the previous report.",
    "Identify the speaker, recorder, time, wording, and affected areas.",
    "Flag the conflict and notify the assistant and affected ministry leaders.",
    "Father, his assistant, or the Ceremony Coordinator identifies the controlling direction.",
    "Notify affected roles and retain the resolved history.",
])

# Future development
doc.add_heading("10. Cross-functional and future ceremony development", level=1)
add_status_box(doc, "PLATFORM REQUIREMENT", "A ceremony is not complete merely because its liturgical portion is documented. Every linked functional area must be reviewed through its designated stewardship.", PALE_BLUE)

add_table(doc, ["Operational case", "Linked workstreams to develop"], [
    ("Confirmation", "Liturgical plan, bishop's retinue, Schola, reception/banquet, venue, food, facilities, setup, cleanup, communications, volunteers."),
    ("Major feast or procession", "Mass and procession, Schola, bearers, banners, hospitality, potluck/banquet, rooms, food, setup, cleanup."),
    ("Annual fundraising gala", "Planning, venue, budget/fundraising, ticketing or RSVP, food, communications, staffing, setup, cleanup, reporting."),
    ("Wedding or Baptism reception", "Family permissions, hall or external venue, access buffers, food responsibility, setup, cleanup, publication."),
], [2700, 6660])

doc.add_heading("Schola release candidate", level=2)
add_bullets(doc, [
    "Document ceremony participation, music selection, singer availability, staffing, notification, and exception handling.",
    "Evaluate the Webmaster's existing Schola capability after the desired workflow is defined.",
    "Include a focused Schola capability in the Scheduling release only if the workflow is approved and reuse meets privacy, maintainability, and integration requirements.",
])

doc.add_heading("Monthly improvement mechanism", level=2)
add_para(doc, "The second-Sunday Monthly Operations Review should examine what worked, what failed, recurring manual interventions, shortages, feedback, and proposed rule changes. Approved rules are documented, effective-dated, tested, and incorporated through controlled change rather than informal memory.")

# Review register
doc.add_heading("11. Stakeholder review register", level=1)
add_table(doc, ["Review item", "Required steward or authority"], [
    ("Exact wedding forms and terminology", "Father / Ceremony Coordinator"),
    ("Funeral rite names, checklists, and Solemn Requiem support", "Father / Ceremony Coordinator"),
    ("External Solemnity and precedence rules", "Father / Ceremony Coordinator"),
    ("Epiphany blessing frequency and times", "Father / Ceremony Coordinator / Stakeholders"),
    ("Additional annual blessings", "Father / Ceremony Coordinator / Stakeholders"),
    ("Schola workflow and release scope", "Schola Workflow Steward / Technical Steward / Product Owner"),
    ("Banquets, potlucks, and Confirmation receptions", "Hospitality/Events Workflow Steward / Facilities Steward"),
    ("Annual fundraising gala", "Events/Fundraising Workflow Steward"),
    ("Final privacy and retention rules", "Product Owner / Privacy and policy review"),
], [5200, 4160])

doc.add_heading("Change control", level=2)
add_para(doc, "When an SME, steward, Father, or stakeholder review changes a rule, record the affected workflow, decision, authority, effective date, permissions, notifications, acceptance criteria, and related manuals or tests. Historical services and prior rules remain auditable.")

doc.add_heading("Authoritative source documents", level=2)
add_bullets(doc, [
    "Generate and Maintain the Regular Service Calendar — Approved July 16, 2026",
    "Schedule a Wedding — Approved for stakeholder review",
    "Schedule a Funeral — Approved for stakeholder review",
    "Schedule a Baptism — Approved for stakeholder review",
    "Select a Requiem Mass for a Scheduled Mass — Approved for stakeholder review",
    "OLV Operations Platform Project Documentation Framework",
    "Chapel Scheduler Architecture and Product Design",
])

add_status_box(doc, "END OF WORKING DRAFT", "This manual should be reviewed at the monthly operations meeting and revised as additional Workflow Stewards complete Schola, hospitality, facilities, fundraising, Confirmation, and other ceremony workstreams.", PALE_GREEN)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT.resolve())
