from pathlib import Path
import sys

sys.path.insert(0, "/private/tmp/chapel-docx")

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/share/Chapel_Scheduler_Project_Roles_Agreement.docx")
BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 90, 90)


def set_font(run, size=11, bold=False, italic=False, color=RGBColor(0, 0, 0)):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn("w:" + edge))
        if node is None:
            node = OxmlElement("w:" + edge)
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.first_child_found_in("w:tblW")
    tblW.set(qn("w:w"), str(sum(widths)))
    tblW.set(qn("w:type"), "dxa")
    tblInd = OxmlElement("w:tblInd")
    tblInd.set(qn("w:w"), "120")
    tblInd.set(qn("w:type"), "dxa")
    tblPr.append(tblInd)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tcW = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tcW.set(qn("w:w"), str(widths[index]))
            tcW.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_labeled(doc, label, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    set_font(p.add_run(label + " "), bold=True, color=DARK)
    set_font(p.add_run(text))
    return p


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.72)
section.bottom_margin = Inches(0.72)
section.left_margin = Inches(0.82)
section.right_margin = Inches(0.82)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 12, 6),
    ("Heading 2", 13, BLUE, 10, 5),
    ("Heading 3", 12, DARK, 8, 4),
):
    style = styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = color
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(header.add_run("CHAPEL SCHEDULER  |  WORKING AGREEMENT"), size=8.5, bold=True, color=MUTED)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(footer.add_run("Chapel-owned project  •  Proposed collaboration framework"), size=8.5, color=MUTED)

kicker = doc.add_paragraph()
kicker.paragraph_format.space_after = Pt(2)
set_font(kicker.add_run("PROJECT GOVERNANCE"), size=9, bold=True, color=BLUE)

title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(3)
set_font(title.add_run("Chapel Scheduler"), size=23, bold=True)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(10)
set_font(subtitle.add_run("Project roles and decision-making agreement"), size=13, color=MUTED)

add_labeled(doc, "Participants:", "Chapel Project Owner and Webmaster / Technical Advisor")
add_labeled(doc, "Purpose:", "Keep product ownership clear while making full use of the webmaster's technical experience, existing components, and integration work.")
add_labeled(doc, "Status:", "Working agreement for discussion; not a legal contract.")

doc.add_heading("Shared objective", level=1)
p = doc.add_paragraph()
set_font(p.add_run("Build a privacy-first, chapel-owned scheduling system that people will actually use. "), bold=True, color=DARK)
set_font(p.add_run("The approved architecture, tested user workflows, and ministry requirements guide implementation. Technology supports those decisions; it does not redefine them by convenience alone."))

doc.add_heading("Roles and decision rights", level=1)
table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
set_table_geometry(table, [2700, 6660])
headers = ("Role", "Primary responsibility")
for i, value in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell._tc.get_or_add_tcPr().append(OxmlElement("w:shd"))
    cell._tc.tcPr[-1].set(qn("w:fill"), "F2F4F7")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_font(p.add_run(value), bold=True, color=DARK)

role_rows = [
    ("Chapel Project Owner", "Owns priorities, requirements, ministry rules, privacy boundaries, UX decisions, acceptance criteria, and final product-scope decisions."),
    ("Webmaster / Technical Advisor", "Recommends the stack; identifies reusable components; advises on security, hosting, integrations, maintainability, effort, and technical risk; implements agreed technical work."),
    ("Implementation contributors", "Build and test against approved requirements, document changes, protect existing data, and raise risks or ambiguities before changing established workflows."),
]
for role, responsibility in role_rows:
    cells = table.add_row().cells
    for i, value in enumerate((role, responsibility)):
        p = cells[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(value), bold=(i == 0), color=DARK if i == 0 else RGBColor(0, 0, 0))
set_table_geometry(table, [2700, 6660])

doc.add_heading("How decisions are made", level=1)
add_labeled(doc, "Product and UX:", "The Chapel Project Owner decides what the system should do and accepts the resulting experience after stakeholder testing.")
add_labeled(doc, "Technical implementation:", "The Webmaster / Technical Advisor recommends how to build it within the approved requirements and explains material tradeoffs in plain English.")
add_labeled(doc, "Safety and feasibility:", "Any contributor may stop a proposed change for a documented security, privacy, legal, reliability, or feasibility concern. The concern is resolved before proceeding.")

doc.add_page_break()
doc.add_heading("Reuse of existing components", level=1)
p = doc.add_paragraph()
set_font(p.add_run("Reuse is encouraged when a component "), bold=True, color=DARK)
set_font(p.add_run("fits the accepted workflow, meets privacy and security requirements, can be maintained by the chapel, and reduces cost or delivery time. A component should be adapted or replaced when using it would force volunteers or administrators into a materially different experience."))

doc.add_heading("Chapel ownership and continuity", level=1)
add_labeled(doc, "Repository and access:", "Source code, documentation, deployment instructions, and issue history live in a chapel-owned organization with at least two chapel-approved owners.")
add_labeled(doc, "Accounts and secrets:", "Production domains, hosting, databases, messaging bots, API credentials, and vendor accounts use chapel-controlled ownership rather than a single person's private account.")
add_labeled(doc, "Handoff readiness:", "Another qualified contributor should be able to operate and maintain the system from the documented repository and controlled credentials.")

doc.add_heading("Working method", level=1)
add_labeled(doc, "Propose:", "Describe the change, user benefit, technical approach, dependencies, and risks.")
add_labeled(doc, "Check:", "Compare it with the architecture, privacy rules, and accepted user workflows.")
add_labeled(doc, "Decide:", "Record the decision and owner before implementation when the change is material.")
add_labeled(doc, "Verify:", "Demonstrate the working result against agreed acceptance criteria before calling it complete.")

doc.add_heading("Review point", level=1)
p = doc.add_paragraph()
set_font(p.add_run("Review this agreement after the initial technical stack is selected and again before production launch. "), bold=True, color=DARK)
set_font(p.add_run("Any revision should preserve chapel ownership, privacy-first design, clear decision rights, and the accepted volunteer experience."))

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT.resolve())
